from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


# Creating a new docx file
document = Document()

# Profile picture
document.add_picture(
    'me.png', width=Inches(2.0)
)

# Details
speak("What is your name")
name = input("What is your name?\n")
speak(f"hello {name} how are you today")
speak("what is your phone number")
phone_number = input("What is your phone number?\n")
speak("what is your email")
email = input("What is you email ?\n")


document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email
)

# about me
document.add_heading('About me')
speak("Tell me about yourself")
about_me = input('Tell me about yourself\n')


# Work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

speak("Enter company name")
company = input("Enter company name\n")
speak(f"From when did you start working at {company}")
from_date = input("From Date\n")
speak(f"Till when did you work at {company}")
to_date = input("To Date\n")


# add_run is used for character level formatting
p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

speak(f"Describe your experience at the {company}")
experience_details = input(
    "Describe your experience at " + company
)


p.add_run(experience_details)

# Adding skills
document.add_heading("Skills")
speak("Enter your skill")
skill = input("Enter your skill\n")
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    speak("Do you have more skills")
    has_more_skills = input("Do you have more skills? (yes/no)\n")

    if has_more_skills.lower() == 'yes':
        speak("enter more skills")
        skill = input("Enter skill\n")

        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generates using python and python-docx package"

# Adding more experiences
while True:
    speak("Do you have more experiences yes or no")
    has_more_experiences = input("Do you have more experiences (yes/no)\n")

    if has_more_experiences.lower() == "yes":
        p = document.add_paragraph()
        speak("Enter your company name")

        company = input("Enter company name\n")
        speak(f"From when did you start working at {company}")
        from_date = input("From Date\n")
        speak(f"Till when did you stat working at{company}")
        to_date = input("To Date\n")


        # add_run is used for character level formatting
        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        speak(f"Describe your experience at {company}")
        experience_details = input(
            "Describe your experience at " + company
        )


        p.add_run(experience_details)
    else:
        break

# Saving document
document.save('cv.docx')
